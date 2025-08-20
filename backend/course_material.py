import pathlib
import os
import re
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# Import LLM helpers
try:
    from llm import get_gemini_client, get_google_search_tool, generate_course_content, system_prompt
except Exception:
    # Allow running without LLM for fallback
    get_gemini_client = None
    get_google_search_tool = None
    generate_course_content = None
    system_prompt = None

def read_all_text_files():
    """Read all text files from the Inputs and Outputs directory"""
    text_files_content = {}
    
    # Read from home directory "Inputs and Outputs"
    home_dir_path = pathlib.Path("Inputs and Outputs")
    
    if not home_dir_path.exists():
        return text_files_content
    
    # Find all .txt files in the directory
    for txt_file in home_dir_path.glob("*.txt"):
        try:
            with open(txt_file, 'r', encoding='utf-8') as f:
                content = f.read()
                text_files_content[txt_file.name] = content
        except Exception as e:
            pass
    
    # Also check copilot subdirectory
    copilot_dir_path = pathlib.Path("copilot/Inputs and Outputs")
    if copilot_dir_path.exists():
        for txt_file in copilot_dir_path.glob("*.txt"):
            try:
                with open(txt_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    # Use relative path as key to avoid conflicts
                    key = f"copilot/{txt_file.name}"
                    text_files_content[key] = content
            except Exception as e:
                pass
    
    return text_files_content

def extract_course_name_from_content(content):
    """Extract course name from the content and format as subject_course"""
    patterns = [
        r'# Course Name:\s*([^\n]+)',  # # Course Name: Course Title
        r'# ([^:\n]+):\s*([^:\n]+)',  # # Main Title: Subtitle  
        r'# ([^:\n]+)',  # # Main Title
        r'## Course Outline:\s*([^\n]+)',  # ## Course Outline: Name
        r'Course Title:\s*([^\n]+)',  # Course Title: Name
        r'Course:\s*([^\n]+)',  # Course: Name
        r'Title:\s*([^\n]+)',  # Title: Name
    ]
    
    for pattern in patterns:
        match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
        if match:
            if len(match.groups()) == 2:  # Title: Subtitle pattern
                course_name = f"{match.group(1).strip()} {match.group(2).strip()}"
            else:
                course_name = match.group(1).strip()
            
            # Extract the main subject and format as subject_course
            course_name = re.sub(r'\s+', ' ', course_name)
            course_name = re.sub(r'[^\w\s-]', '', course_name)
            course_name = course_name.strip()
            
            # Extract key subject words and format
            words = course_name.lower().split()
            key_words = []
            for word in words:
                if word not in ['advanced', 'introduction', 'to', 'and', 'the', 'of', 'in', 'course', 'fundamentals', 'basics']:
                    key_words.append(word)
            
            if key_words:
                subject_name = '_'.join(key_words[:3])  # Take first 3 key words
                return f"{subject_name}_course"
    
    return "course_material"

def parse_weeks_from_content(content):
    """Parse the content to extract individual weeks"""
    weeks = []
    
    # Look for week headers and find content between them
    week_headers = list(re.finditer(r'# Week (\d+):', content, re.IGNORECASE))
    
    if week_headers:
        for i, match in enumerate(week_headers):
            week_num = int(match.group(1))
            start_pos = match.start()
            
            # Find content until next week or end
            if i + 1 < len(week_headers):
                end_pos = week_headers[i + 1].start()
                week_content = content[start_pos:end_pos]
            else:
                week_content = content[start_pos:]
            
            weeks.append({
                'number': week_num,
                'type': 'Week',
                'content': week_content.strip(),
                'title': f"Week {week_num}"
            })
    else:
        # Fallback: treat entire content as one section
        weeks.append({
            'number': 1,
            'type': 'Course',
            'content': content.strip(),
            'title': "Complete Course Content"
        })
    
    weeks.sort(key=lambda x: x['number'])
    return weeks

def _extract_section(content: str, header_regex: str) -> str | None:
    """Extract section text by header regex until next top-level header or end."""
    match = re.search(header_regex, content, re.IGNORECASE)
    if not match:
        return None
    start = match.end()
    next_match = re.search(r"\n##?\s+|\n#\s+Week\s+\d+", content[start:], re.IGNORECASE)
    if next_match:
        end = start + next_match.start()
    else:
        end = len(content)
    return content[start:end].strip()


def _parse_weeks_simple(content: str):
    """Parse '# Week N:' blocks with their content."""
    weeks = []
    for m in re.finditer(r"^#\s*Week\s*(\d+)\s*:(.*)$", content, re.MULTILINE | re.IGNORECASE):
        num = int(m.group(1))
        title_line = m.group(0)
        start = m.end()
        next_m = re.search(r"^#\s*Week\s*\d+\s*:.*$", content[start:], re.MULTILINE | re.IGNORECASE)
        end = start + next_m.start() if next_m else len(content)
        block = content[start:end].strip()
        weeks.append({
            'number': num,
            'title': re.sub(r'^#\s*', '', title_line).strip(),
            'content': block,
        })
    weeks.sort(key=lambda x: x['number'])
    return weeks


def sanitize_filename(name: str) -> str:
    name = name.strip().replace('\n', ' ')
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    name = re.sub(r"\s+", " ", name)
    return name or "course_material"


def create_combined_docx(content, course_title, output_dir):
    """Create DOCX with layout: Title, Course Overview, Weekly Summary, Week sections (each new page)."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading(f'{course_title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Course Overview
    doc.add_heading('Course Overview', 1)
    doc.add_paragraph('─' * 60)
    overview_text = _extract_section(content, r"##\s*(Course Overview|Overview|Course Description)\s*")
    doc.add_paragraph(overview_text or "")

    # Weekly Summary
    weekly_summary = _extract_section(content, r"##\s*Weekly\s*Summary\s*")
    if weekly_summary:
        doc.add_paragraph()
        doc.add_heading('Weekly Summary', 1)
        doc.add_paragraph('─' * 60)
        for line in weekly_summary.split('\n'):
            t = line.strip()
            if not t:
                continue
            t = re.sub(r'^[-*]\s*', '', t)
            doc.add_paragraph(t, style='List Bullet')

    # Detailed Week sections: each on a new page
    weeks = _parse_weeks_simple(content)
    for week_idx, week in enumerate(weeks):
        if week_idx > 0:
            doc.add_page_break()
        # Add week heading
        week_title_line = week['title'].replace('#', '').strip()
        doc.add_heading(week_title_line, 1)
        doc.add_paragraph('─' * 60)
        
        # Render week content paragraphs and lists
        lines = week['content'].split('\n')
        current_paragraph = ""
        in_code_block = False
        for line in lines:
            original_line = line
            line = line.strip()
            if not line:
                if current_paragraph:
                    if '**' in current_paragraph:
                        p = doc.add_paragraph()
                        parts = current_paragraph.split('**')
                        for i, part in enumerate(parts):
                            run = p.add_run(part)
                            if i % 2 == 1:
                                run.bold = True
                    else:
                        doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_paragraph()
                continue
            if line.startswith('```'):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                in_code_block = not in_code_block
                continue
            if in_code_block:
                p = doc.add_paragraph(original_line)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Inches(0.1)
                continue
            if line.startswith('#### '):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_heading(line.replace('#### ', ''), 4)
            elif line.startswith('### '):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_heading(line.replace('### ', ''), 3)
            elif line.startswith('## '):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_heading(line.replace('## ', ''), 2)
            elif line.startswith('- ') or line.startswith('* '):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                bullet_text = line.replace('- ', '').replace('* ', '')
                if '**' in bullet_text:
                    p = doc.add_paragraph(style='List Bullet')
                    parts = bullet_text.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    doc.add_paragraph(bullet_text, style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                list_text = re.sub(r'^\d+\. ', '', line)
                if '**' in list_text:
                    p = doc.add_paragraph(style='List Number')
                    parts = list_text.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    doc.add_paragraph(list_text, style='List Number')
            elif line.startswith('---') or line == '=' * len(line):
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_paragraph('─' * 60)
            else:
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
        if current_paragraph:
            if '**' in current_paragraph:
                p = doc.add_paragraph()
                parts = current_paragraph.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                doc.add_paragraph(current_paragraph)
        
    # Save document
    filename = f"{sanitize_filename(course_title)}.docx"
    filepath = os.path.join(output_dir, filename)
    
    try:
        doc.save(filepath)
        return filepath
    except Exception as e:
        return None

def create_combined_pdf(content, course_title, output_dir):
    """Create PDF with layout: Title, Course Overview, Weekly Summary, Week sections (page breaks)."""
    filename = f"{sanitize_filename(course_title)}.pdf"
    filepath = os.path.join(output_dir, filename)
    
    try:
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=1*inch,
            leftMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor='darkblue'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor='blue'
        )
        
        week_heading_style = ParagraphStyle(
            'WeekHeading',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=15,
            spaceBefore=20,
            textColor='darkblue'
        )
        
        section_heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=15,
            textColor='darkblue'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
            leftIndent=0,
            rightIndent=0
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=4,
            leftIndent=20,
            bulletIndent=10
        )
        
        elements = []
        
        # Add title
        elements.append(Paragraph(course_title, title_style))
        elements.append(Spacer(1, 30))
        
        # Course Overview
        elements.append(Paragraph("Course Overview", week_heading_style))
        elements.append(Spacer(1, 10))
        
        overview_text = _extract_section(content, r"##\s*(Course Overview|Overview|Course Description)\s*") or ""
        if overview_text:
            try:
                elements.append(Paragraph(overview_text, normal_style))
            except:
                clean_overview = re.sub(r'<[^>]+>', '', overview_text)
                elements.append(Paragraph(clean_overview, normal_style))
        
        elements.append(Spacer(1, 20))

        # Weekly Summary
        weekly_summary = _extract_section(content, r"##\s*Weekly\s*Summary\s*")
        if weekly_summary:
            elements.append(Paragraph("Weekly Summary", week_heading_style))
            elements.append(Spacer(1, 6))
            for line in weekly_summary.split('\n'):
                t = line.strip()
                if not t:
                    continue
                t = re.sub(r'^[-*]\s*', '', t)
                try:
                    elements.append(Paragraph(f"• {t}", bullet_style))
                except:
                    clean_b = re.sub(r'<[^>]+>', '', t)
                    elements.append(Paragraph(f"• {clean_b}", bullet_style))
            elements.append(Spacer(1, 20))
        
        # Detailed week content - each week starts on a new page
        weeks = _parse_weeks_simple(content)
        if weeks:
            for week_idx, week in enumerate(weeks):
                # Add week heading
                week_title = week['title'].replace('#', '').strip()
                
                elements.append(Paragraph(week_title, week_heading_style))
                elements.append(Spacer(1, 10))
                
                # Process week content
                lines = week['content'].split('\n')
                current_paragraph = ""
                
                for line in lines:
                    line = line.strip()
                    
                    if not line:
                        if current_paragraph:
                            # Clean and format paragraph
                            clean_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', current_paragraph)
                            clean_para = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean_para)
                            try:
                                elements.append(Paragraph(clean_para, normal_style))
                            except:
                                clean_para = re.sub(r'<[^>]+>', '', clean_para)
                                elements.append(Paragraph(clean_para, normal_style))
                            current_paragraph = ""
                        elements.append(Spacer(1, 6))
                        continue
                    
                    # Skip the main week header (already processed)
                    if line.startswith('# Week '):
                        continue
                    
                    # Handle headings
                    if line.startswith('### ') or line.startswith('## '):
                        if current_paragraph:
                            clean_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', current_paragraph)
                            try:
                                elements.append(Paragraph(clean_para, normal_style))
                            except:
                                clean_para = re.sub(r'<[^>]+>', '', clean_para)
                                elements.append(Paragraph(clean_para, normal_style))
                            current_paragraph = ""
                        
                        heading_text = line.replace('### ', '').replace('## ', '')
                        elements.append(Paragraph(heading_text, section_heading_style))
                        elements.append(Spacer(1, 5))
                    
                    # Handle bullet points
                    elif line.startswith('- ') or line.startswith('* '):
                        if current_paragraph:
                            clean_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', current_paragraph)
                            try:
                                elements.append(Paragraph(clean_para, normal_style))
                            except:
                                clean_para = re.sub(r'<[^>]+>', '', clean_para)
                                elements.append(Paragraph(clean_para, normal_style))
                            current_paragraph = ""
                        
                        bullet_text = line.replace('- ', '').replace('* ', '')
                        bullet_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', bullet_text)
                        try:
                            elements.append(Paragraph(f"• {bullet_text}", bullet_style))
                        except:
                            clean_bullet = re.sub(r'<[^>]+>', '', bullet_text)
                            elements.append(Paragraph(f"• {clean_bullet}", bullet_style))
                    
                    # Handle separator lines
                    elif line.startswith('---') or line == '=' * len(line):
                        if current_paragraph:
                            clean_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', current_paragraph)
                            try:
                                elements.append(Paragraph(clean_para, normal_style))
                            except:
                                clean_para = re.sub(r'<[^>]+>', '', clean_para)
                                elements.append(Paragraph(clean_para, normal_style))
                            current_paragraph = ""
                        elements.append(Spacer(1, 10))
                    
                    # Regular content
                    else:
                        if current_paragraph:
                            current_paragraph += " " + line
                        else:
                            current_paragraph = line
                
                # Add any remaining paragraph
                if current_paragraph:
                    clean_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', current_paragraph)
                    try:
                        elements.append(Paragraph(clean_para, normal_style))
                    except:
                        clean_para = re.sub(r'<[^>]+>', '', clean_para)
                        elements.append(Paragraph(clean_para, normal_style))
                
                # Add page break between weeks (except for last week)
                if week_idx < len(weeks) - 1:
                    elements.append(PageBreak())
        
        # Build PDF
        doc.build(elements)
        return filepath
        
    except Exception as e:
        return None

def build_structured_text_llm(raw_corpus: str, planner_text: str, title_hint: str | None) -> str | None:
    """Use LLM to produce strict structured text with required layout."""
    if not (get_gemini_client and get_google_search_tool and generate_course_content and system_prompt):
        return None
    client = get_gemini_client()
    tool = get_google_search_tool()
    task = (
        "ROLE & GOAL\n"
        "You are an expert instructional designer. Produce a clean, strictly structured plain text/Markdown syllabus ready for DOCX/PDF conversion. \n"
        "Incorporate and prioritize the user's planner input. Maintain clarity, coherence, and completeness.\n\n"
        "AUDIENCE & STYLE\n"
        "- Audience: match the level and style signaled in the planner (e.g., Intermediate, Conceptual & Conversational).\n"
        "- Tone: clear, structured, conversational-but-professional. Define jargon on first use.\n"
        "- Be self-contained; avoid references to prompts, agents, or external tools.\n\n"
        "STRICT TOP-LEVEL STRUCTURE (must follow exactly)\n"
        "1) Title line (plain text, not a header): the exact course name from planner.\n"
        "2) '## Course Overview' – a single concise paragraph.\n"
        "3) '## Weekly Summary' – one bullet per week (1–2 lines each).\n"
        "4) Week sections ONLY, each starting with a top-level header: '# Week N: <short title>' in ascending order.\n"
        "   After the last week, STOP. Do NOT add any other top-level sections.\n\n"
        "WEEK SECTION REQUIREMENTS (expand substantially)\n"
        "- Target length per week: ~500–800 words.\n"
        "- Within each week, include the following labeled subsections (plain text labels, not Markdown headers):\n"
        "  Concepts: 3–6 concise bullets defining and explaining key ideas.\n"
        "  Example: a concrete, step-by-step illustrative example.\n"
        "  Case Study: a realistic scenario (3–6 sentences) connecting ideas to practice.\n"
        "  Exercise: one brief activity or reflection question.\n"
        "  Tip/Pitfall: a common misconception or a tip to remember.\n"
        "- Use short paragraphs and bullets for readability.\n"
        "- Ensure week content is explanatory and builds intuition.\n\n"
        "CONSTRAINTS & GUARDRAILS\n"
        "- Do NOT include any of the following: prerequisites, resources, references, grading, schedules, agent/system prompts, or meta-instructions.\n"
        "- Do NOT include Markdown tables or images. Use only plain text and simple bullets.\n"
        "- Keep numbering correct for weeks (N starts at 1).\n"
        "- Do not hallucinate facts about the user's institution; keep examples generic and plausible.\n\n"
        "QUALITY CHECK BEFORE FINALIZING\n"
        "- Title matches planner.\n"
        "- Overview is 1 paragraph.\n"
        "- Weekly Summary has exactly one bullet per week.\n"
        "- Each week has all required labeled subsections and sufficient elaboration.\n"
        "- No extra sections after the last week.\n"
    )
    # Build course_content payload combining planner and corpus
    course_content = (
        (f"PLANNER INPUT:\n{planner_text}\n\n" if planner_text else "") +
        (f"OTHER TEXT INPUTS:\n{raw_corpus}" if raw_corpus else "")
    )
    try:
        response = generate_course_content(
            client=client,
            teaching_style="",
            duration="",
            difficulty_level="",
            google_search_tool=tool,
            system_prompt=system_prompt,
            filepath=pathlib.Path("Inputs and Outputs/curriculum.pdf"),
            course_content=course_content,
            task=task,
        )
        return response.text
    except Exception:
        return None


def extract_title_from_planner(text: str) -> str | None:
    # Try explicit course name patterns first
    m = re.search(r"^#\s*Course Name:\s*(.+)$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    # Fall back to first H1 header
    m = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    # Or bolded course name mention
    m = re.search(r"\*\*Course Name\*\*\s*[:\-]?\s*(.+)$", text, re.IGNORECASE | re.MULTILINE)
    if m:
        return m.group(1).strip()
    return None

def write_structured_txt(structured_text: str, course_name: str, output_dir: str) -> str | None:
    """Write the structured TXT to disk and return its path."""
    try:
        txt_path = os.path.join(output_dir, f"{sanitize_filename(course_name)}.txt")
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(structured_text)
        return txt_path
    except Exception:
        return None

def get_duration_weeks():
    """Read `user_config.json` and return the number of weeks as int if available."""
    try:
        with open('user_config.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
        duration = data.get('duration', '')
        match = re.search(r'(\d+)\s*week', duration, re.IGNORECASE)
        if match:
            return int(match.group(1))
    except Exception:
        pass
    return None

def should_only_list_weeks(content: str) -> bool:
    """Return True if there are no explicit week sections in content and we have a duration."""
    has_week_sections = re.search(r'#\s*Week\s*\d+\s*[:\-]?', content or '', re.IGNORECASE) is not None
    duration_weeks = get_duration_weeks()
    return (not has_week_sections) and bool(duration_weeks and duration_weeks > 0)

def get_subject_from_config() -> str | None:
    """Return a sanitized subject string from user_config.json if available."""
    try:
        with open('user_config.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
        # Possible keys to look for
        raw = data.get('subject') or data.get('course_subject') or data.get('course_name')
        if raw and isinstance(raw, str):
            # Sanitize to snake_case words
            cleaned = re.sub(r'[^\w\s-]', '', raw).strip().lower()
            cleaned = re.sub(r'\s+', ' ', cleaned)
            parts = [p for p in cleaned.split(' ') if p]
            if parts:
                return '_'.join(parts)
    except Exception:
        pass
    return None

def main():
    """Main function to generate combined course materials"""
    # Read all text files
    text_files = read_all_text_files()
    
    # Combine all content
    combined_content = ""
    
    # Prioritize enhanced content if available
    if 'enhanced_course_content.txt' in text_files:
        combined_content = text_files['enhanced_course_content.txt']
    else:
        # Combine all available content
        for filename, content in text_files.items():
            combined_content += f"\n\n=== CONTENT FROM {filename.upper()} ===\n"
            combined_content += content
            combined_content += f"\n=== END OF {filename.upper()} ===\n"
    
    # Clean content - remove teaching agent prompts and system messages
    lines = combined_content.split('\n')
    cleaned_lines = []
    skip_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Skip lines that contain teaching agent prompts or system messages
        if any(keyword in line_lower for keyword in [
            'teaching agent', 'system prompt', 'agent prompt', 'instruction:', 
            'you are a', 'act as', 'your role', 'generate', 'create a course',
            'llm', 'ai assistant', 'chatgpt', 'copilot'
        ]):
            skip_section = True
            continue
            
        # Skip empty lines after filtering
        if skip_section and line.strip() == '':
            continue
        elif line.strip() != '':
            skip_section = False
            
        # Keep the line if it's not filtered
        if not skip_section:
            cleaned_lines.append(line)
    
    combined_content = '\n'.join(cleaned_lines)
    
    # Read planner agent instruction for title and as LLM input
    planner_path = pathlib.Path("Inputs and Outputs/planner_agent_instruction.txt")
    planner_text = ""
    if planner_path.exists():
        try:
            planner_text = planner_path.read_text(encoding='utf-8')
        except Exception:
            planner_text = planner_path.read_text(errors='ignore')

    # Determine course title from planner
    course_title = extract_title_from_planner(planner_text) or extract_course_name_from_content(planner_text or combined_content)
    # If still generic, try config subject
    if not course_title or course_title == 'course_material':
        subject = get_subject_from_config()
        if subject:
            course_title = subject.replace('_', ' ').title()
        else:
            course_title = "Course"
    
    # Create output directory named "course material" inside "Inputs and Outputs"
    output_dir = os.path.join("Inputs and Outputs", "course material")
    
    try:
        os.makedirs(output_dir, exist_ok=True)
    except Exception as e:
        print(f"Error creating directory: {e}")
        return
    
    # Build structured TXT via LLM first; fallback to simple assembly if needed
    structured_text = build_structured_text_llm(combined_content, planner_text, course_title)
    if not structured_text:
        # Minimal fallback ensuring required headers exist
        lines = [f"# Course Name: {course_title}", "", "## Course Overview", "", "## Weekly Summary", "- Week 1", "", "# Week 1: Introduction", "Content TBD"]
        structured_text = "\n".join(lines)
    txt_path = write_structured_txt(structured_text, course_name=course_title, output_dir=output_dir)
    
    # Create combined DOCX and PDF from structured text
    docx_path = create_combined_docx(structured_text, course_title, output_dir)
    pdf_path = create_combined_pdf(structured_text, course_title, output_dir)
    
    # Only show success message if both files created
    if docx_path and pdf_path:
        print("Course materials created successfully!")
        base = sanitize_filename(course_title)
        print(f"Files created: {base}.txt, {base}.docx and {base}.pdf")
    else:
        print("Error creating course materials")

if __name__ == "__main__":
    main()
